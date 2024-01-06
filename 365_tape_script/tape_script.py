import docx
from datetime import date

# Create date variables.
today = date.today()
d1 = today.strftime("%m/%d/%Y")
d2 = today.strftime("%m-%d-%Y")

# convert date to string for saving document
string_date = str(d2)

# Print Introduction and Version number
print("Hello! Welcome to Tape Scanner v0.1")
print("Please scan Box and Tape barcodes below.")

# Ask user to continue
while True:
    print("Press 'ENTER' to Continue.")
    user = input()
    if user:
        break

    # Create a document and add a paragraph to the document.
    doc = docx.Document()
    p = doc.add_paragraph()

    # Add a Title run to the paragraph and format
    title_run = p.add_run(d1 + " - DOR-CO.AURORA.103-NBU SCRATCH PROCESSING-DAILY")
    title_run.bold = True
    title_run.font.name = 'Calibri (Body)'
    title_run.font.size = docx.shared.Pt(11)

    # Add another pargraph
    p = doc.add_paragraph()

    # Add a Section run and format it
    section_run = p.add_run("Please update the library inventory. \n")
    section_run = p.add_run("Scratch tapes loaded today: \n \n")
    section_run.font.name = 'Calibri (Body)'
    section_run.font.size = docx.shared.Pt(11)

    # Add another Box run
    box_run = p.add_run("Box #: \n")
    box_run.bold = True

    # Accept user input for box number
    box_number = input("Scan BOX number: ")
    box = p.add_run(box_number + "\n\n")

    # Another Tape run
    tape_run = p.add_run("Tapes: \n")
    tape_run.bold = True

    # Accept multiple user input for tape numbers
    while True:
        tape_number = input("Scan TAPES (Or input 'Q' to Quit and Save after scanning the last one): ")
        if tape_number == 'Q':
            # Save the document
            doc.save('Tapes ' + string_date + '.docx')
            break
        else:
            # Add run to document
            tapes = p.add_run(tape_number + "\n")
    break